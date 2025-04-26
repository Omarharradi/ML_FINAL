
# utils.py

import numpy as np
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
import streamlit as st  
import plotly.express as px
import os
import base64




def get_filtered_df(df, selected_dashboards, selected_positions, selected_individuals, selected_type):
    filtered = df.copy()
    if selected_dashboards:
        filtered = filtered[filtered["# Dashboard"].isin(selected_dashboards)]
    if selected_positions:
        filtered = filtered[filtered["Position"].isin(selected_positions)]
    if selected_individuals:
        filtered = filtered[filtered["Leader"].isin(selected_individuals)]
    if selected_type:
        filtered = filtered[filtered["Typology 1"].isin(selected_type)]
    return filtered


def build_donut_chart(lis_data):
    ''' 
    LIS Stands for Leadership Index Score which is a weighted score of critical skills necessary skills and beneficial skills to have a standardised metric to compare all individuals
    EQ assesses Emotional Intelligence which is one of the most important skills that all leaders are assessed in.
    '''
    mean_lis = np.mean(lis_data)
    std_lis = np.std(lis_data)
    std_low = 70
    std_high = 84

    leaders_meeting = np.sum((lis_data >= std_low) & (lis_data <= std_high))
    leaders_exceeding = np.sum(lis_data > std_high)
    leaders_requiring_training = np.sum(lis_data < std_low)
    
    leaders_summary = {
    "meeting_standard": np.sum((lis_data >= std_low) & (lis_data < std_high)),
    "exceeding_standard": np.sum(lis_data >= std_high),
    "requiring_training": np.sum(lis_data < std_low)
    }

       # Original full data
    full_labels = ['Demonstrating Role <br>Proficiency', 'Surpassing Role Proficiency', 'Requiring Training']
    full_values = [leaders_meeting, leaders_exceeding, leaders_requiring_training]
    full_colors = ['#5c9acc', '#f4a300', '#e63946']

    # Filter out 0% categories
    filtered_labels = []
    filtered_values = []
    filtered_colors = []
    for label, value, color in zip(full_labels, full_values, full_colors):
        if value > 0:
            filtered_labels.append(label)
            filtered_values.append(value)
            filtered_colors.append(color)

    fig = go.Figure(data=[go.Pie(
        labels=filtered_labels,
        values=filtered_values,
        hole=0.4,
        marker=dict(colors=filtered_colors),
        hovertemplate="%{label}: %{value} leaders (%{percent})",
        textinfo='percent+label',
        textposition='outside',
        rotation=100
    )])

    fig.update_layout(
        title="Overall Leadership Competency Levels",
        margin=dict(t=80, b=80, l=50, r=50),
        width=700,
        height=400,
        showlegend=False  # <-- this line removes the legend
    )

    return fig, leaders_summary


def build_histogram(lis_data):
    ''' 
    LIS Stands for Leadership Index Score which is a weighted score of critical skills necessary skills and beneficial skills to have a standardised metric to compare all individuals
    EQ assesses Emotional Intelligence which is one of the most important skills that all leaders are assessed in.
    '''
    mean_lis = np.mean(lis_data)
    std_lis = np.std(lis_data)
    std_low = mean_lis - 1.5 * std_lis
    std_high = mean_lis + 1.5 * std_lis

    below_std = np.sum(lis_data < std_low)
    above_std = np.sum(lis_data > std_high)

    # Summary dictionary
    summary_stats = {
        "mean": mean_lis,
        "std_dev": std_lis,
        "1.5_std_below": std_low,
        "1.5_std_above": std_high,
        "min_value": np.min(lis_data),
        "max_value": np.max(lis_data),
        "count": len(lis_data),
        "below_1.5_std": below_std,
        "above_1.5_std": above_std
    }
    

    fig = go.Figure()
    fig.add_trace(go.Histogram(
        x=lis_data,
        nbinsx=20,
        marker_color='darkblue',
        opacity=0.75,
        name='LIS'
    ))
    fig.add_vline(x=mean_lis, line=dict(color='gold', dash='dash'),
                  annotation_text=f'Mean: {mean_lis:.2f}', annotation_position="top right")
    fig.add_vline(x=std_low, line=dict(color='green', dash='dash'),
                  annotation_text=f'1.5-std below: {std_low:.2f}', annotation_position="top left")
    fig.add_vline(x=std_high, line=dict(color='green', dash='dash'),
                  annotation_text=f'1.5-std above: {std_high:.2f}', annotation_position="top right")
    fig.update_layout(
        title="Leadership Index Score (LIS) Distribution",
        xaxis_title="Leadership Index Score (LIS) Score",
        yaxis_title="Frequency",
        template="plotly_white",
        width=700,
        height=400,
        margin=dict(t=50, b=50, l=50, r=50)
    )
    return fig, summary_stats

import numpy as np
import pandas as pd
import plotly.graph_objects as go
from plotly.subplots import make_subplots
def radar_chart_plotly(dashboard, leader, df, skills_mapping):
    '''
    Compare average dashboard skill scores and individual leader's skill scores.
    If leader is "None", only show the dashboard average scores.
    '''
    mapping = skills_mapping.get(dashboard, {})
    categories = ['Critical Skills', 'Necessary', 'Beneficial Skills']

    summary_stats = {
        "dashboard": dashboard,
        "leader": leader,
        "skill_scores": {}
    }

    color_palette = px.colors.qualitative.Pastel
    category_colors = {
        "Average": color_palette[0],
        "Leader": color_palette[1]
    }

    fig = make_subplots(
        rows=1, cols=len(categories),
        specs=[[{'type': 'polar'} for _ in categories]],
        subplot_titles=[f"{cat}" for cat in categories],
    )

    radial_range = [0, 100]
    tickvals = [0, 40, 80]

    for i, cat in enumerate(categories):
        skills = mapping.get(cat, [])
        if not skills:
            continue

        # Average dashboard scores
        avg_scores = [df.loc[df['# Dashboard'] == dashboard, skill].mean() for skill in skills if skill in df.columns]
        
        # If a leader is selected, get their scores; otherwise, set leader scores to None
        if leader != "None":
            leader_scores = [df.loc[(df['# Dashboard'] == dashboard) & (df['Leader'] == leader), skill].values[0]
                             if skill in df.columns else None for skill in skills]
        else:
            leader_scores = [None] * len(skills)  # No leader selected, show no leader scores

        summary_stats["skill_scores"][cat] = {
            "average": avg_scores,
            "leader": leader_scores
        }

        # Format skill names for radial labels
        formatted_skills = [
            skill.replace(" ", "<br>").replace("-", "<br>") if "&" not in skill
            else skill.replace(" &", "&").replace("-", "<br>").replace(" ", "<br>").replace("&", " &")
            for skill in skills
        ]
        skills_closed = formatted_skills + [formatted_skills[0]]

        # Close the loop for radar plot
        avg_scores_closed = np.concatenate((avg_scores, [avg_scores[0]]))
        leader_scores_closed = np.concatenate((leader_scores, [leader_scores[0]]))

        # Add average trace
        fig.add_trace(
            go.Scatterpolar(
                r=avg_scores_closed,
                theta=skills_closed,
                fill='toself',
                mode='lines+markers',
                name=f"Dashboard Avg",
                marker=dict(color=category_colors["Average"])
            ),
            row=1, col=i+1
        )

        # Only add the leader trace if a leader is selected
        if leader != "None":
            fig.add_trace(
                go.Scatterpolar(
                    r=leader_scores_closed,
                    theta=skills_closed,
                    fill='toself',
                    mode='lines+markers',
                    name=f"{leader} (Leader)",
                    marker=dict(color=category_colors["Leader"])
                ),
                row=1, col=i+1
            )

        # Update radar settings
        fig.update_polars(
            dict(
                radialaxis=dict(visible=True, range=radial_range, tickvals=tickvals, showline=True),
                angularaxis=dict(
    rotation=50,
    tickfont=dict(size=10),
    direction="clockwise"
)
            ),
            row=1, col=i+1
        )

    fig.update_layout(
        title_text=f"Radar Chart: {leader} vs Dashboard Avg ({dashboard})" if leader != "None" else f"Radar Chart: {dashboard} (No Leader)",
        showlegend=True,
        width=350 * len(categories),
        height=400,
        margin=dict(t=80, b=0, l=60, r=100),
        legend=dict(
    orientation="h",        # horizontal legend
    yanchor="bottom",
    y=-0.2,
    xanchor="center",
    x=0.5,
    title_text='',          # remove default legend title
    font=dict(size=10)
)
    )

    return fig, summary_stats



import plotly.express as px

def build_polar_chart(df, metric="LIS"):
    """
    Builds a polar bar chart for the given metric ('LIS' or 'EQ') by Typology.

    Args:
        df (pd.DataFrame): Data with 'Typology 1', 'LIS', and 'EQ' columns
        metric (str): Either "LIS" or "EQ"

    Returns:
        fig (plotly.graph_objs._figure.Figure): Polar bar chart
        avg_stats (list of dict): Summary per typology
    """
    assert metric in ["LIS", "EQ"], "Metric must be 'LIS' or 'EQ'"

    avg = df.groupby("Typology 1")[metric].mean().reset_index()
    avg_stats = avg.to_dict(orient="records")

    fig = px.bar_polar(
        avg,
        r=metric,
        theta="Typology 1",
        color="Typology 1",
        template="plotly_white",
        title=f"Average {metric} by Typology",
        color_discrete_sequence=px.colors.qualitative.Pastel
    )

    fig.update_layout(
        margin=dict(t=80, b=40, l=40, r=40),
        polar=dict(
            radialaxis=dict(visible=True, range=[0, 100])
        )
    )

    return fig, avg_stats


def build_box_plot(df):
    ''' 
    LIS Stands for Leadership Index Score which is a weighted score of critical skills necessary skills and beneficial skills to have a standardised metric to compare all individuals
    EQ assesses Emotional Intelligence which is one of the most important skills that all leaders are assessed in.
    '''

    summary_stats = df.groupby("Dashboard Number")["Overall Results"].describe().reset_index()

    fig = px.box(
        df,
        x="Dashboard Number",
        y="Overall Results",
        title="Box Plot of EQ by Dashboard",
        template="plotly_white",
        color="Dashboard Number",
        color_discrete_sequence=px.colors.qualitative.Pastel
    )
    return fig, summary_stats

def get_insights_chart(lis_data, source_code, llm):
    insights_template = PromptTemplate(
        input_variables=["Dataframe", "PlotCode"],
        template="""You are a data analyst helping decision-makers understand data trends.

    Step 1: Analyze the PlotCode and DataFrame to understand what the chart is visualizing. Do NOT include or describe the code itself.

    Step 2: Return your output in **Markdown** using the following structure:

    **Chart Description**
    \n
    A brief paragraph that explains what the chart is visualizing and why it's relevant.
    **Chart Insights**
    - Use concise bullet points to highlight key patterns, trends, or outliers in the data.
    - Make the language natural and insightful, not robotic.
    
    Important Guidelines:
    - If a KPI, variable, or concept is not clearly defined in the DataFrame or PlotCode, do NOT guess its meaning. It may represent proprietary or internal data.
    - Never fabricate interpretations — only base your insights on what is explicitly observable from the data and plot logic.
    - Be clear, thoughtful, and avoid assumptions.

    Only return the final Markdown-formatted insight. The tone should be for non-technical audience, make it simple.

    DataFrame:
    {Dataframe}

    PlotCode:
    {PlotCode}
    """
    )

    #llm = ChatGoogleGenerativeAI(model="gemini-2.0-flash")
    chain = insights_template | llm | StrOutputParser()
    response = chain.invoke({"Dataframe": lis_data, "PlotCode": source_code})
    return response



# utils.py


from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_core.prompts import PromptTemplate
from langchain_experimental.agents import create_pandas_dataframe_agent
from langchain_core.runnables import RunnableSequence
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import RetrievalQA
from langchain.memory import ConversationBufferMemory  
from docx import Document as DocxDocument
from langchain.embeddings import HuggingFaceEmbeddings
import io
from contextlib import redirect_stdout
from langchain_experimental.tools.python.tool import PythonAstREPLTool
from langchain.agents import Tool, AgentType, initialize_agent




def initialize_pipeline(docx_path="N&P LDP Final Report Strategy & Enablement Plan.docx"):
    """
    Initializes the pipeline by:
      - Loading Word doc and converting it into Langchain Documents.
      - Creating a vector store from it.
      - Initializing pandas agent from df.
      - Setting up classifier and answer reranking.
    """

    # -----------------------------
    # Load & Process Word Document
    # -----------------------------
    raw_text = ""
    doc = DocxDocument(docx_path)
    for para in doc.paragraphs:
        if para.text.strip():
            raw_text += para.text.strip() + "\n"

# Split the raw text into manageable chunks for embedding
    text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1500,
    chunk_overlap=200
)
    chunks = text_splitter.split_text(raw_text)
    docs = [Document(page_content=chunk) for chunk in chunks]  # try with first 30 only
    # -----------------------------
    # Load DataFrame from CSV
    # -----------------------------
    df = pd.read_csv('LDP_summary.csv')
    df = df[['ID', 'Leader', 'Position', 'LIS', 'EQ', 'Typology 1',
             '# Dashboard', 'Skills_Below_Threshold', 'Should_be_promoted',
             'Engagement Score', 'Dashboard Number']]

    # -----------------------------
    # LLM & Embeddings
    # -----------------------------
    llm = ChatGoogleGenerativeAI(model="gemini-1.5-pro-latest")
    embedding = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-mpnet-base-v2"
)
    # -----------------------------
    # RAG Setup
    # -----------------------------
    persist_dir = os.path.abspath("./chroma_db")
    db_file = os.path.join(persist_dir, "chroma.sqlite3")

    # Ensure the directory exists
    os.makedirs(persist_dir, exist_ok=True)

    # Load or create the vector store
    if os.path.exists(db_file):
        print("✅ Loading existing Chroma vector store...")

        vectorstore = Chroma(
            persist_directory="./chroma_db",
            embedding_function=embedding,
            collection_name='ldp_docs'
        )
    else:
        print("🆕 Creating new Chroma vector store...")

        vectorstore = Chroma.from_documents(
            documents=docs,
            embedding=embedding,
            collection_name='ldp_docs',
            persist_directory="./chroma_db",
        )

        vectorstore.persist()
        print("💾 Chroma DB persisted to disk.")
    
    retriever = vectorstore.as_retriever(search_kwargs={"k": 20})

    # Custom prompt for the RAG chain
    context_prompt = PromptTemplate.from_template("""
You are a data assistant helping analyze leadership assessment data.

Use the provided context to answer the question below. The context includes information on individuals, dashboards, and their skills.

Context:
{context}

Question:
{question}

Answer by combining insights across individuals. Focus on patterns, averages, and summaries. If no useful context is found, say "Not enough information. Never mention the context."
""")

    # Using StuffDocumentsChain to combine documents into context
    from langchain.chains.combine_documents.stuff import StuffDocumentsChain as CombineStuffDocumentsChain
    stuff_chain = create_stuff_documents_chain(
    llm=llm,
    prompt=context_prompt,
    document_variable_name="context"
)
    
    memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)


    rag_chain = RetrievalQA.from_chain_type(
    llm=llm,
    chain_type="stuff",  # Explicitly specify the chain type
    retriever=retriever,
    input_key="query",
    memory=memory
)

    # -----------------------------
    # Pandas Agent Setup
    # -----------------------------

    # Initialize PythonAstREPLTool with df
    python_tool = PythonAstREPLTool(locals={"df": df})

# Define tools with improved description
    tools = [
    Tool.from_function(
        name="Python DataFrame Tool",
        func=python_tool.run,
        description="Executes Python code to query the DataFrame `df`. Input must be valid Python code (e.g., `df.columns.tolist()` or `df[df['is_leader'] == True].shape[0]`). Use this to analyze or retrieve data from `df`. Do not redefine or simulate `df`."
    )
]

    # Initialize agent with improved prefix
    pandas_agent = initialize_agent(
    tools=tools,
    llm=llm,
    agent=AgentType.ZERO_SHOT_REACT_DESCRIPTION,
    verbose=True,
    handle_parsing_errors=True,
    agent_kwargs={
            "prefix": """
You are a smart data analyst assistant. You are working with a pandas DataFrame called `df`, which contains the results of a leadership development program.

Each row in the DataFrame represents one unique individual (a leader).

Here is an overview of the DataFrame structure:
<class 'pandas.core.frame.DataFrame'>
RangeIndex: 85 entries, 0 to 84
Data columns (total 11 columns):
 0   ID                      → Unique identifier for each leader
 1   Leader                  → Full name of the leader (string)
 2   Position                → The job title or current role of the leader
 3   LIS                     → Leadership Index Score (float). Ranges from 0 to 100. Higher means stronger leadership fit.
 4   EQ                      → Emotional Intelligence score (float). Ranges from 0 to 100.
 5   Typology 1              → Personality-based leadership typology (e.g., "The Mentoring Leader")
 6   # Dashboard             → Dashboard group name and context
 7   Skills_Below_Threshold  → Number of skill areas in which the leader scored below a critical threshold
 8   Should_be_promoted      → Boolean indicating whether the individual is recommended for promotion
 9   Engagement Score        → An engagement metric score (float, may have missing values)
10  Dashboard Number         → Code for the dashboard group (e.g., D1, D2)

General rules and instructions:
- Do NOT create or assume the existence of new columns (e.g., no 'is_leader' or 'score_level').
- Use only the actual columns listed above. You can always check them via `df.columns.tolist()` if needed.
- Never simulate or redefine `df`. You are querying a real DataFrame already loaded.
- Use proper Python code to analyze the DataFrame. Examples:
    - To count total leaders: `len(df)`
    - To find the average LIS: `df['LIS'].mean()`
    - To get top performers: `df.sort_values('LIS', ascending=False).head(5)`
    - To filter by condition: `df[df['LIS'] > 85]`

Always reason step-by-step, and answer based only on what the DataFrame actually contains.
If a user’s request is ambiguous, first examine the column names and structure before deciding how to proceed.
"""
}
)


    # -----------------------------
    # Classifier to Route Query
    # -----------------------------
    classifier_prompt = PromptTemplate.from_template("""
You are a smart classifier. Given a user question, decide if it should be handled using:

- "structured" → if it's about rankings, filters, math, comparisons, scores, means, who should be promoted, etc. If the question is asking about who has the highest score, or if a specific leader is below a threshold, it's structured.
- "semantic" → if it's about summaries, meaning, types of skills, descriptions. Anything that has to do with insights is semantic as well. Basically, if the question is not about who (which leader, position or dashboard) it is most likely semantic. Any other question that is not about a specific leader, position or dashboard or about getting statistics like the mean is semantic.

Return ONLY the word "structured" or "semantic".

Question: {query}
""")
    classifier_chain = classifier_prompt | llm

    # -----------------------------
    # Unified Ask Function
    # -----------------------------
    def ask(query: str) -> str:
        route_output = classifier_chain.invoke({"query": query})
        route = route_output.content.strip().lower()
        print(f"[Routing → {route}]")
        
        if route == "structured":
            result = pandas_agent.run(query)
            raw_result = f"Returned:\n{result}"

            print("Raw result:", raw_result)            
            # Use an LLM chain to rephrase the full chain trace into a natural response.
            rephrase_prompt = PromptTemplate.from_template("""
You are a data analyst assistant. Below is a trace from a data analysis agent including thoughts, code, and printed results.

Trace:
{raw_result}

The user asked:
"{query}"

Your job is to generate an accurate and complete answer based only on the actual printed numbers in the trace.
- Ignore any 'Final Answer:' lines if they contradict the correct printed results.
- Do not make assumptions or guess values.
- Use names, values, and stats exactly as shown above.
- Do not summarize or simplify unless the trace clearly does.
- Do not talk about the trace or the agent.

Write a clean and clear response.
""")
            print("Raw result:", rephrase_prompt)
            rephrase_chain = rephrase_prompt | llm
            natural_result = rephrase_chain.invoke({"raw_result": raw_result, "query": query})
            
            return natural_result.content
        elif route == "semantic":
            response = rag_chain.invoke({"query": query})
            return str(response["result"]).strip()
        else:
            return "Sorry, I couldn't confidently classify your question."

    return ask



import inspect
import streamlit as st

def display_insight(unique_key, build_func, lis_data, llm, expander_title=None):
    """
    Generates and displays insights for a given chart using session state.

    Parameters:
    - unique_key (str): A unique identifier for the chart (e.g., "donut", "hist", etc.).
    - build_func (function): The chart function whose source code will be used.
    - lis_data: Data to be passed to the get_insights_chart function.
    - llm: The LLM instance to generate insights.
    - expander_title (str, optional): Title for the expander. Defaults to "Insights for {unique_key}".
    """
    # Ensure that the session state dictionary exists
    if "insights" not in st.session_state:
        st.session_state["insights"] = {}

    # Generate insight when the button is pressed
    if st.button("🧠 Explain this graph", key=f"explain_{unique_key}"):
        source_code = inspect.getsource(build_func)
        with st.spinner("Generating insights..."):
            response = get_insights_chart(lis_data, source_code, llm)
        st.session_state["insights"][unique_key] = response

    # Display the saved insight, if it exists
    if unique_key in st.session_state["insights"]:
        title = expander_title if expander_title else f"Insights for {unique_key}"
        with st.expander(title, expanded=True):
            st.info(st.session_state["insights"][unique_key])


def get_image_base64(image_path):
    with open(image_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode()


def dynamic_sidebar_filters(df):
    """Render dynamic sidebar filters and return the filtered DataFrame + current selections."""
    with st.sidebar:
        # --- Logo section with base64 HTML ---
        nesma_base64 = get_image_base64("assets/nesma.png")
        ivy_base64 = get_image_base64("assets/Ivy Logo copy.png")

        st.markdown(f"""
            <div style="padding-top: 10px; padding-bottom: -5px;">
                <div style="display: flex; justify-content: start; align-items: center; gap: 10px;">
                    <img src="data:image/png;base64,{nesma_base64}" alt="Nesma Logo" width="120" style="vertical-align: middle;">
                    <img src="data:image/png;base64,{ivy_base64}" alt="Ivy Logo" width="120" style="vertical-align: middle;">
                </div>
            </div>
        """, unsafe_allow_html=True)

        st.markdown("---")

        st.header("Filters")

        # 1. Dashboard filter (no dependencies at this point)
        all_dashboards = sorted(df["# Dashboard"].dropna().unique())
        selected_dashboards = st.multiselect("Select Dashboard(s)", all_dashboards, key="Select Dashboard(s)", help="Select one or more dashboards to filter the data.")

        # 2. Filter DataFrame based on selected dashboards (if any)
        df_dash_filtered = df[df["# Dashboard"].isin(selected_dashboards)] if selected_dashboards else df

        # 3. Position filter based on dashboards
        position_options = sorted(df_dash_filtered["Position"].dropna().unique()) if "Position" in df.columns else []
        selected_positions = st.multiselect("Select Position(s)", position_options, help="Select one or more positions to filter the data.", key="Select Position(s)")

        # 4. Further filter DataFrame based on positions
        df_pos_filtered = df_dash_filtered[df_dash_filtered["Position"].isin(selected_positions)] if selected_positions else df_dash_filtered

        # 5. Individual filter based on dashboards + positions
        individual_options = sorted(df_pos_filtered["Leader"].dropna().unique()) if "Leader" in df.columns else []
        selected_individuals = st.multiselect("Select Leader(s)", individual_options, help="Select one or more leaders to filter the data.", key="Select Leader(s)")

        df_type_filtered = df_dash_filtered[df_dash_filtered["Leader"].isin(selected_individuals)] if selected_individuals else df_dash_filtered

        # 5. Individual filter based on dashboards + positions
        type_options = sorted(df_type_filtered["Typology 1"].dropna().unique()) if "Typology 1" in df.columns else []
        selected_type = st.multiselect("Select Leadership Type(s)", type_options, help="Select one or more typology to filter the data.", key="Select Type(s)")

        # Checkbox to toggle "All Leaders selected" state
        all_selected = not any([selected_dashboards, selected_positions, selected_individuals, selected_type])
        toggle_all = st.checkbox("All Leaders selected", value=all_selected)

        # If checkbox is unchecked by user (i.e. filters applied, but then unticked), reset all selections
        if toggle_all and not all_selected:
            selected_dashboards.clear()
            selected_positions.clear()
            selected_individuals.clear()
            selected_type.clear()

    # --- Final filtering: apply all selected filters together ---
    df_filtered = df.copy()

    if selected_dashboards:
        df_filtered = df_filtered[df_filtered["# Dashboard"].isin(selected_dashboards)]
    if selected_positions:
        df_filtered = df_filtered[df_filtered["Position"].isin(selected_positions)]
    if selected_individuals:
        df_filtered = df_filtered[df_filtered["Leader"].isin(selected_individuals)]
    if selected_type:
        df_filtered = df_filtered[df_filtered["Typology 1"].isin(selected_type)]

    return df_filtered, selected_dashboards, selected_positions, selected_individuals, selected_type





# insight_charts.py


def plot_resource_type_distribution(df):
    """
    Plots a stacked bar chart showing the distribution of resource types per skill category.

    Parameters:
    - df (pd.DataFrame): DataFrame containing 'Skill Category' and 'Resource Type' columns.

    Returns:
    - fig (plotly.graph_objs._figure.Figure): Plotly figure object.
    """
    resource_type_by_category = df.groupby(['Skill Category', 'Resource Type']).size().reset_index(name='Count')
    fig = px.bar(resource_type_by_category, x='Skill Category', y='Count', color='Resource Type',
                 title='Distribution of Resource Types per Skill Category', barmode='stack')
    return fig, resource_type_by_category

def plot_top_skills(df, top_n=10):
    """
    Plots a bar chart of the top N most frequently targeted skills.

    Parameters:
    - df (pd.DataFrame): DataFrame containing a 'Skill' column.
    - top_n (int): Number of top skills to display (default is 10).

    Returns:
    - fig (plotly.graph_objs._figure.Figure): Plotly figure object.
    """
    skill_leader_counts = (
        df.groupby('Skill')['Leader']
        .nunique()
        .sort_values(ascending=False)
        .head(top_n)
        .reset_index()
        .rename(columns={'Leader': 'Leader Count'})
    )

    fig = px.bar(
        skill_leader_counts,
        x='Skill',
        y='Leader Count',
        title=f'Top {top_n} Skills by Number of Unique Leaders Assessed',
        labels={'Leader Count': 'Number of Leaders'},
    )
    return fig, skill_leader_counts




def plot_resource_distribution_by_category(df):
    """
    Plots a pie chart showing the distribution of total resources per skill category.

    Parameters:
    - df (pd.DataFrame): DataFrame containing 'Skill Category' and 'Resource Text' columns.

    Returns:
    - fig (plotly.graph_objs._figure.Figure): Plotly figure object.
    """
    skill_cat_resources = df.groupby('Skill Category')['Resource Text'].count().reset_index(name='Total Resources')
    fig = px.pie(skill_cat_resources, names='Skill Category', values='Total Resources',
                 title='Resource Distribution by Skill Category')
    return fig, skill_cat_resources



import plotly.express as px

def plot_top_resources_by_category(df, category, top_n=8):
    """
    Plots a vertical bar chart of the top N most frequent resources within a given category.

    Parameters:
    - df (pd.DataFrame): DataFrame with at least 'Resource Type' and 'Resource Text' columns.
    - category (str): The resource category to filter on.
    - top_n (int): The number of top resources to show.

    Returns:
    - fig (plotly.graph_objs._figure.Figure): Vertical bar chart of top resources.
    - top_resources (pd.DataFrame): DataFrame of top resource texts and their counts.
    """
    # Filter by selected resource category
    filtered_df = df[df['Resource Type'] == category]

    if filtered_df.empty:
        print(f"No resources found for category: {category}")
        return None, None

    # Count top resources
    top_resources = (
        filtered_df['Resource Text']
        .value_counts()
        .head(top_n)
        .reset_index()
    )
    top_resources.columns = ['Resource Text', 'Count']
    top_resources['Resource Text'] = top_resources['Resource Text'].apply(
    lambda x: '<br>'.join(x[i:i+20] for i in range(0, len(x), 20))
)

    # Create vertical bar chart
    fig = px.bar(
        top_resources,
        x='Resource Text',
        y='Count',
        title=f"Top {top_n} Most Recurrent Resources in '{category}'",
        text_auto=True,
        height=500
    )

    fig.update_layout(
    xaxis_tickangle=20,
   xaxis_tickfont=dict(size=10)
)
    return fig, top_resources



def plot_leaders_below_threshold(group, top_n=10, selected_skill=None):
    """
    Plots a vertical bar chart of the top N leaders with the most skills below threshold.

    Parameters:
    - group (pd.DataFrame): DataFrame containing 'Leader', 'Below_Threshold_Count', and 'Skill'.
    - top_n (int): Number of top leaders to display (default is 10).
    - selected_skill (str, optional): Specific skill to filter by. If None, shows all skills.

    Returns:
    - fig (plotly.graph_objs._figure.Figure): Plotly figure object.
    """
    
    # Filter by the selected skill (if provided)
    if selected_skill:
        filtered_group = group[group['Skill'] == selected_skill]
    else:
        filtered_group = group

    # Get leaders with skills below threshold
    leaders_below = filtered_group[filtered_group['Below_Threshold_Count'] == True]['Leader'].value_counts().reset_index()
    leaders_below.columns = ['Leader', 'Below Threshold Count']

    # Get the top N leaders with the most skills below threshold
    leaders_below = leaders_below.head(top_n)

    # Create the vertical bar chart
    fig = px.bar(leaders_below, x='Leader', y='Below Threshold Count',
                 title=f'Top {top_n} Leaders with the Most Skills Below Threshold',
                 color='Below Threshold Count')
    
    # Update layout to make the chart vertical
    fig.update_layout(
        xaxis_title='Leader',
        yaxis_title='Below Threshold Count',
        xaxis={'tickangle': 45},  # Rotate x-axis labels for better readability
        height=500
    )
    
    return fig, leaders_below




import pandas as pd
import plotly.express as px

def plot_top_performers_by_skill(df, skill=None, percentile=0.9, bar_color="#636EFA"):
    """
    Plots a clean bar chart showing the top X% performers for a selected skill, including tied scores.

    Parameters:
    - df (pd.DataFrame): DataFrame containing at least 'Leader', 'Skill', and 'Score' columns.
    - skill (str or None): Specific skill to filter by. If None, no plot is returned.
    - percentile (float): Percentile threshold (e.g., 0.9 for top 10%) (default is 0.9).
    - bar_color (str): Hex color code for the bars (default is Plotly's gray-blue).

    Returns:
    - fig (plotly.graph_objs._figure.Figure): Plotly figure object.
    - top_performers_df (pd.DataFrame): DataFrame of top performers.
    """
    if not skill:
        print("Please select a skill.")
        return None, pd.DataFrame()

    filtered_df = df[df['Skill'] == skill]

    if filtered_df.empty:
        print(f"No data found for skill: {skill}")
        return None, pd.DataFrame()

    # Max score per leader to avoid duplicates
    unique_leader_scores = filtered_df.groupby('Leader', as_index=False)['Score'].max()

    # Compute cutoff based on percentile
    threshold = unique_leader_scores['Score'].quantile(percentile)

    # Filter top performers (including ties)
    top_performers_df = unique_leader_scores[unique_leader_scores['Score'] >= threshold].sort_values(by='Score', ascending=False)

    fig = px.bar(
        top_performers_df,
        x='Leader',
        y='Score',
        title=f"Top 10% Performers in '{skill}'",
        text='Score'
    )

    fig.update_traces(marker_color=bar_color, textposition='outside')
    fig.update_layout(
        yaxis_title='Score',
        xaxis_title='Leader',
        uniformtext_minsize=8,
        uniformtext_mode='hide',
        plot_bgcolor='white',
        title_x=0.5,
        margin=dict(t=60, b=40),
    )

    return fig, top_performers_df






def build_histogram_with_leaders_eq(df, highlight_leaders=None):
    '''
    Expects:
        - df with columns: 'EQ' and 'Leader'
        - highlight_leaders (optional): list of leader names to highlight
    '''
    lis_data = df['EQ']
    leaders = df['Leader']
    
    # Stats
    mean_lis = np.mean(lis_data)
    std_lis = np.std(lis_data)
    std_low = mean_lis - 1.5 * std_lis
    std_high = mean_lis + 1.5 * std_lis

    # Bin the LIS data
    bin_counts, bin_edges = np.histogram(lis_data, bins=45)
    bin_centers = 0.5 * (bin_edges[:-1] + bin_edges[1:])

    bin_indices = np.digitize(lis_data, bin_edges, right=False) - 1
    bin_indices = np.clip(bin_indices, 0, len(bin_centers) - 1)
    df['EQ_bins'] = bin_centers[bin_indices]

    # Group leaders per bin
    grouped = df.groupby('EQ_bins').agg({
        'Leader': lambda x: ', '.join(sorted(set(x))),
        'EQ': 'count'
    }).reset_index().rename(columns={'EQ': 'Count'})

    # Default bar color
    bar_colors = ['darkblue'] * len(grouped)

    # Highlight bins if leaders selected
    if highlight_leaders:
        highlight_leaders = set([l.strip().lower() for l in highlight_leaders])
        for i, center in enumerate(grouped['EQ_bins']):
            leaders_in_bin = set(map(str.lower, df[df['EQ_bins'] == center]['Leader']))
            if highlight_leaders & leaders_in_bin:
                bar_colors[i] = 'crimson'

    # Build the plot
    fig = go.Figure()

    fig.add_trace(go.Bar(
        x=grouped['EQ_bins'],
        y=grouped['Count'],
        customdata=grouped['Leader'],
        marker_color=bar_colors,
        name='EQ Distribution',
        hovertemplate='<b>EQ Center:</b> %{x}<br>' +
                      '<b>Count:</b> %{y}<br>' +
                      '<b>Leaders:</b> %{customdata}<extra></extra>'
    ))

    EPS = 1e-9                       # small positive number
    std_safe = std_lis if std_lis > 0 else EPS

    # Bell curve
    x_vals = np.linspace(min(lis_data), max(lis_data), 500)
    y_vals = (
        (1 / (std_safe * np.sqrt(2 * np.pi))) *
        np.exp(-0.5 * ((x_vals - mean_lis) / std_safe) ** 2)
    )
    y_scaled = y_vals * max(grouped['Count']) / max(y_vals)

    fig.add_trace(go.Scatter(
        x=x_vals,
        y=y_scaled,
        mode='lines',
        name='Normal Curve',
        line=dict(color='orange', dash='dash')
    ))

    # Vertical lines
    for val, label, color in zip([mean_lis, std_low, std_high],
                                 ['Mean', '-1.5 Std', '+1.5 Std'],
                                 ['green', 'red', 'red']):
        fig.add_vline(
            x=val,
            line=dict(color=color, dash='dot'),
            annotation_text=label,
            annotation_position="top",
            annotation_font_color=color
        )

    fig.update_layout(
        title="EQ Distribution with Highlight",
        xaxis_title="EQ Score",
        yaxis_title="Number of People",
        template="plotly_white",
        width=850,
        height=500,
        margin=dict(t=60, b=50, l=60, r=40),
        legend=dict(yanchor="top", y=0.98, xanchor="left", x=0.01)
    )


        # Summary stats
    summary_stats = {
        "mean": mean_lis,
        "std_dev": std_lis,
        "1.5_std_below": std_low,
        "1.5_std_above": std_high,
        "min_value": np.min(lis_data),
        "max_value": np.max(lis_data),
        "count": len(lis_data),
        "below_1.5_std": np.sum(lis_data < std_low),
        "above_1.5_std": np.sum(lis_data > std_high)
    }

    return fig, summary_stats


def build_histogram_with_leaders(df, highlight_leaders=None):
    '''
    Expects:
        - df with columns: 'LIS' and 'Leader'
        - highlight_leaders (optional): list of leader names to highlight
    '''
    lis_data = df['LIS']
    leaders = df['Leader']
    
    # Stats
    mean_lis = np.mean(lis_data)
    std_lis = np.std(lis_data)
    std_low = mean_lis - 1.5 * std_lis
    std_high = mean_lis + 1.5 * std_lis

    # Bin the LIS data
    bin_counts, bin_edges = np.histogram(lis_data, bins=20)
    bin_centers = 0.5 * (bin_edges[:-1] + bin_edges[1:])

    bin_indices = np.digitize(lis_data, bin_edges, right=False) - 1
    bin_indices = np.clip(bin_indices, 0, len(bin_centers) - 1)
    df['LIS_bin_center'] = bin_centers[bin_indices]

    # Group leaders per bin
    grouped = df.groupby('LIS_bin_center').agg({
        'Leader': lambda x: ', '.join(sorted(set(x))),
        'LIS': 'count'
    }).reset_index().rename(columns={'LIS': 'Count'})

    # Default bar color
    bar_colors = ['darkblue'] * len(grouped)

    # Highlight bins if leaders selected
    if highlight_leaders:
        highlight_leaders = set([l.strip().lower() for l in highlight_leaders])
        for i, center in enumerate(grouped['LIS_bin_center']):
            leaders_in_bin = set(map(str.lower, df[df['LIS_bin_center'] == center]['Leader']))
            if highlight_leaders & leaders_in_bin:
                bar_colors[i] = 'crimson'

    # Build the plot
    fig = go.Figure()

    fig.add_trace(go.Bar(
        x=grouped['LIS_bin_center'],
        y=grouped['Count'],
        customdata=grouped['Leader'],
        marker_color=bar_colors,
        name='LIS Distribution',
        hovertemplate='<b>LIS Center:</b> %{x}<br>' +
                      '<b>Count:</b> %{y}<br>' +
                      '<b>Leaders:</b> %{customdata}<extra></extra>'
    ))
    EPS = 1e-9                       # small positive number
    std_safe = std_lis if std_lis > 0 else EPS

    # Bell curve
    x_vals = np.linspace(min(lis_data), max(lis_data), 500)
    y_vals = (
        (1 / (std_safe * np.sqrt(2 * np.pi))) *
        np.exp(-0.5 * ((x_vals - mean_lis) / std_safe) ** 2)
    )
    y_scaled = y_vals * max(grouped['Count']) / max(y_vals)

    fig.add_trace(go.Scatter(
        x=x_vals,
        y=y_scaled,
        mode='lines',
        name='Normal Curve',
        line=dict(color='orange', dash='dash')
    ))

    # Vertical lines
    for val, label, color in zip([70, 85],
                                 ['Below 70', '85+'],
                                 ['red', 'green']):
        fig.add_vline(
            x=val,
            line=dict(color=color, dash='dot'),
            annotation_text=label,
            annotation_position="top",
            annotation_font_color=color
        )

    fig.update_layout(
        title="Leadership Index Score (LIS) Distribution with Highlight",
        xaxis_title="LIS Score",
        yaxis_title="Number of People",
        template="plotly_white",
        width=850,
        height=500,
        margin=dict(t=60, b=50, l=60, r=40),
        legend=dict(yanchor="top", y=0.98, xanchor="left", x=0.01)
    )


        # Summary stats
    summary_stats = {
        "mean": mean_lis,
        "std_dev": std_lis,
        "1.5_std_below": std_low,
        "1.5_std_above": std_high,
        "min_value": np.min(lis_data),
        "max_value": np.max(lis_data),
        "count": len(lis_data),
        "below_1.5_std": np.sum(lis_data < std_low),
        "above_1.5_std": np.sum(lis_data > std_high)
    }

    return fig, summary_stats



def plot_typology_distribution(df):
    """
    Plot the distribution of 'Typology 1' as a pie chart without using red.

    Parameters:
    - df: pd.DataFrame, the input dataframe
    """
    if 'Typology 1' in df.columns:
        typology_counts = df['Typology 1'].value_counts().reset_index()
        typology_counts.columns = ['Typology', 'Count']

        # Custom colors (excluding red)
        custom_colors = [
    "#1F77B4",  # deep blue
    "#2CA02C",  # vivid green
    "#9467BD",  # royal purple
    "#FF7F0E",  # bold orange
    "#17BECF",  # bright cyan
    "#3CB371",  # medium sea-green
    "#7F7F7F",  # charcoal gray
    "#BCBD22"   # olive-gold
]

        typologies = typology_counts['Typology'].unique()
        color_map = {typ: custom_colors[i % len(custom_colors)] for i, typ in enumerate(typologies)}

        fig_typ_dist = px.pie(
            typology_counts,
            names='Typology',
            values='Count',
            title='Leadership Typology Distribution',
            color='Typology',
            color_discrete_map=color_map
        )

        return fig_typ_dist, typology_counts
    





def plot_strongest_and_weakest_skills(filtered_df, top_n=7):
    """
    Generates two bar plots:
    - Top N Strongest Skills by Average Score
    - Top N Weakest Skills by Average Score

    Parameters:
    - filtered_df (pd.DataFrame): DataFrame with 'Skill' and 'Score' columns.
    - top_n (int): Number of top skills to display in each chart.

    Returns:
    - fig_strong (plotly.graph_objs._figure.Figure): Plot of strongest skills.
    - fig_weak (plotly.graph_objs._figure.Figure): Plot of weakest skills.
    - strongest_df (pd.DataFrame): DataFrame of strongest skills and average scores.
    - weakest_df (pd.DataFrame): DataFrame of weakest skills and average scores.
    """
    
    # Strongest skills by average score
    strongest_df = (
        filtered_df.groupby("Skill")["Score"]
        .mean()
        .sort_values(ascending=False)
        .head(top_n)
        .reset_index()
    )

    fig_strong = px.bar(
        strongest_df,
        x="Skill",
        y="Score",
        title=f"Top {top_n} Strongest Skills by Average Score",
        labels={"Score": "Average Score"},
        text_auto=True
    )

    fig_strong.update_layout(width=800, height=500)


    # Weakest skills by average score
    weakest_df = (
        filtered_df.groupby("Skill")["Score"]
        .mean()
        .sort_values()
        .head(top_n)
        .reset_index()
    )

    fig_weak = px.bar(
        weakest_df,
        x="Skill",
        y="Score",
        title=f"Top {top_n} Weakest Skills by Average Score",
        labels={"Score": "Average Score"},
        text_auto=True
    )
    fig_weak.update_layout(width=800, height=530)
    summary_stats = {
    "strongest": strongest_df,   # top-N highest averages
    "weakest":   weakest_df      # top-N lowest averages
}


    return fig_strong, fig_weak, summary_stats



import plotly.express as px

def plot_training_buckets_per_skill(filtered_df):
    """
    Generates a stacked bar plot for the training buckets per skill and proficiency level,
    with sorting and custom colors for the proficiency levels.

    Parameters:
    - filtered_df (pd.DataFrame): DataFrame containing 'Skill', 'Level', and 'DASH ID' columns.

    Returns:
    - fig (plotly.graph_objs._figure.Figure): The stacked bar plot figure.
    - summary_stat (pd.DataFrame): Summary of the count of training buckets per skill and proficiency level.
    """
    # Calculate counts of skills per level
    skill_level_counts = (
        filtered_df.groupby(['Skill', 'Level'])
        .size()
        .reset_index(name='Count')
    )

    # Sort the skill_level_counts by the 'Count' column in descending order
    skill_level_counts = skill_level_counts.sort_values(by='Count', ascending=False)

    # Define a custom color scale with 'Beginner' as red
    color_scale = {
    'Beginner (0-59)': '#6BAED6',       # Light blue
    'Intermediate (60-74)': '#3182BD',  # Medium blue
    'Advanced (75-100)': '#08519C'      # Dark blue
}
    # Create the bar plot
    fig = px.bar(
        skill_level_counts,
        x="Skill",
        y="Count",
        color="Level",
        title="Training Buckets per Skill and Proficiency Level",
        barmode="stack",
        color_discrete_map=color_scale
    )
    fig.update_layout(xaxis_tickangle=45, height=600)

    return fig, skill_level_counts



def plot_recommended_resources_by_skill(personal_data, selected_leader):
    """
    Plots the number of recommended resources per skill for a given leader.

    Parameters:
    - personal_data (pd.DataFrame): DataFrame containing 'Skill' and 'Resource Text' columns.
    - selected_leader (str): The name of the leader to display in the title.

    Returns:
    - fig (plotly.graph_objs._figure.Figure): Bar plot of resource counts per skill.
    - resource_count (pd.DataFrame): DataFrame of skills and their corresponding resource count.
    """
    # Count resources per skill
    resource_count = (
        personal_data.groupby("Skill")["Resource Text"]
        .count()
        .reset_index()
        .rename(columns={"Resource Text": "Resource Count"})
        .sort_values(by="Resource Count", ascending=False)
    )

    # Create bar plot
    fig = px.bar(
        resource_count,
        x="Skill",
        y="Resource Count",
        title=f"Recommended Resources by Skill for {selected_leader}",
        text_auto=True
    )

    return fig, resource_count




def plot_eq_leader_skills(beginner=4, intermediate=48, advanced=30):
    """
    Plots a bar chart showing the number of leaders by EQ skill levels:
      - Beginner: Scored 0-59% On Skill
      - Intermediate: Scored 60-74% On Skill
      - Advanced: Scored 75-84% On Skill

    Parameters:
        beginner (int): Number of leaders at the beginner level.
        intermediate (int): Number of leaders at the intermediate level.
        advanced (int): Number of leaders at the advanced level.
    """
    # Define the categories and counts
    categories = [
        "Beginner (Scored 0-59%)",
        "Intermediate (Scored 60-74%)",
        "Advanced (Scored 75-84%)"
    ]
    counts = [beginner, intermediate, advanced]

    # Create the bar plot using Plotly
    fig = go.Figure(
        data=go.Bar(
            x=categories,
            y=counts,
            marker_color='indigo',  # Customize color as needed
            text=counts,  # Display the count on top of each bar
            textposition='outside'  # Automatically positions the text above the bars
        )
    )

    # Update layout for better readability
    fig.update_layout(
        title="EQ Leader Skill Levels",
        xaxis_title="Skill Level",
        yaxis_title="Number of Leaders",
        template="plotly_white"
    )

    # Display the figure
    return fig



def build_typology_bar_chart(df, metric="LIS", highlight_leaders=None):
    """
    Builds a bar chart for the given metric ('LIS' or 'EQ') by Typology.

    Args:
        df (pd.DataFrame): Data with 'Typology 1', 'LIS', and 'EQ' columns
        metric (str): Either "LIS" or "EQ"
        highlight_typologies (list, optional): Typologies to highlight

    Returns:
        fig (plotly.graph_objs._figure.Figure): Bar chart
        avg_stats (list of dict): Summary per typology
    """
    assert metric in ["LIS", "EQ"], "Metric must be 'LIS' or 'EQ'"

    # Compute average metric per typology
    avg = df.groupby("Typology 1")[metric].mean().reset_index()
    avg_stats = avg.to_dict(orient="records")

    # Default color for bars
    bar_colors = ['darkblue'] * len(avg)

    if highlight_leaders:
        # Normalize to lowercase for comparison
        highlight_leaders = set(l.lower() for l in highlight_leaders)

        # Find which typologies have selected leaders
        typologies_with_highlights = df[df['Leader'].str.lower().isin(highlight_leaders)]["Typology 1"].unique()

        # Highlight those bars
        for i, typ in enumerate(avg["Typology 1"]):
            if typ in typologies_with_highlights:
                bar_colors[i] = 'crimson'

    # Create the bar plot
    fig = go.Figure()

    fig.add_trace(go.Bar(
        x=avg["Typology 1"],
        y=avg[metric],
        marker_color=bar_colors,
        name=f"Average {metric}",
        hovertemplate="<b>%{x}</b><br>Avg " + metric + ": %{y:.2f}<extra></extra>"
    ))

    fig.update_layout(
        title=f"Average {metric} by Typology",
        xaxis_title="Typology",
        yaxis_title=f"{metric} Score",
        template="plotly_white",
        margin=dict(t=60, b=60, l=60, r=40),
        height=500
    )

    return fig, avg_stats


import plotly.graph_objects as go
import pandas as pd

def plot_lis_by_typology(df):
    """
    Plots a bar chart showing average LIS score by Typology.

    Args:
        df (pd.DataFrame): Must contain 'Typology 1' and 'LIS' columns.

    Returns:
        fig (plotly.graph_objs.Figure): The bar chart figure.
    """
    # Group by Typology and calculate average LIS
    avg_lis = df.groupby("Typology 1")["LIS"].mean().reset_index()

    # Build bar chart
    fig = go.Figure(go.Bar(
        x=avg_lis["Typology 1"],
        y=avg_lis["LIS"],
        marker_color='steelblue',
        hovertemplate="<b>%{x}</b><br>Avg LIS: %{y:.2f}<extra></extra>"
    ))

    # Layout customization
    fig.update_layout(
        title="Average LIS Score by Typology",
        xaxis_title="Typology",
        yaxis_title="Average LIS Score",
        template="plotly_white",
        margin=dict(t=60, b=60, l=60, r=40),
        height=500
    )

    return fig
